import html
from pathlib import Path
from typing import Callable

import streamlit as st
import yaml
from jinja2 import Template
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate

PROFILE_FILE = Path("profile.yaml")
DEGREES_FILE = Path("degrees.yaml")
CERTIFICATIONS_FILE = Path("certifications.yaml")
LANGUAGE_OPTIONS = {
    "en": {
        "label": "English",
        "template_path": Path("template.txt"),
        "skill_suffix": "",
    },
    "de": {
        "label": "Deutsch",
        "template_path": Path("template_de.txt"),
        "skill_suffix": "_de",
    },
}
DEFAULT_LANGUAGE = "en"
SKILL_BLOCKS = [
    {"id": "intro", "label_key": "section_intro", "context_key": "skills_intro", "file_stem": "skills_intro"},
    {
        "id": "market_dev",
        "label_key": "section_job_a",
        "context_key": "skills_job_block_a",
        "file_stem": "skills_job_block_a",
    },
    {
        "id": "senior_bi",
        "label_key": "section_job_b",
        "context_key": "skills_job_block_b",
        "file_stem": "skills_job_block_b",
    },
    {
        "id": "project_manager",
        "label_key": "section_job_c",
        "context_key": "skills_job_block_c",
        "file_stem": "skills_job_block_c",
    },
    {
        "id": "business_analytics",
        "label_key": "section_job_d",
        "context_key": "skills_job_block_d",
        "file_stem": "skills_job_block_d",
    },
    {"id": "fpna", "label_key": "section_job_e", "context_key": "skills_job_block_e", "file_stem": "skills_job_block_e"},
]
LANGUAGE_STRINGS = {
    "en": {
        "language_select_label": "Language / Sprache",
        "candidate_expander": "Candidate Details",
        "full_name_label": "Full Name",
        "full_name_placeholder": "Full Name",
        "address_label": "Address",
        "address_placeholder": "Address",
        "city_label": "City",
        "city_placeholder": "City",
        "region_label": "Region",
        "region_placeholder": "Region/State",
        "country_label": "Country",
        "country_placeholder": "Country",
        "company_expander": "Company Details",
        "company_name_label": "Company Name",
        "company_name_placeholder": "Company Name",
        "position_label": "Job Title / Position",
        "position_placeholder": "Job Title",
        "hiring_manager_label": "Hiring Manager",
        "hiring_manager_placeholder": "Hiring Manager",
        "job_city_label": "City",
        "job_city_placeholder": "City",
        "job_region_label": "Region",
        "job_region_placeholder": "Region/State",
        "job_country_label": "Country",
        "job_country_placeholder": "Country",
        "awareness_label": "How did you hear about the role?",
        "awareness_placeholder": "a LinkedIn post",
        "sections_expander": "Text Blocks",
        "section_candidate_block": "Candidate block",
        "section_company_block": "Company block",
        "section_subject_block": "Subject block",
        "section_greeting_block": "Greeting block",
        "section_intro": "Intro & skills",
        "section_job_a": "Job block A",
        "section_job_b": "Job block B",
        "section_job_c": "Job block C",
        "section_job_d": "Job block D",
        "section_job_e": "Job block E",
        "section_degrees": "Degrees & Certifications",
        "section_languages": "Languages & why this company",
        "section_goodbye": "Goodbye block",
        "qualifications_expander": "Qualifications & Certifications",
        "qualifications_caption": "Choose which degrees and certifications appear or add custom entries.",
        "degrees_heading": "Degrees",
        "degrees_input_label": "Add custom degree",
        "add_degree_placeholder": "e.g., MSc Big Data & Analytics",
        "add_degree_button": "Add degree",
        "select_all_degrees": "Select all degrees",
        "clear_degrees": "Clear degrees",
        "degrees_multiselect_label": "Degrees to include",
        "degrees_multiselect_help": "These entries feed the 'I hold ...' sentence.",
        "certifications_heading": "Certifications",
        "certs_input_label": "Add custom certification",
        "add_cert_placeholder": "e.g., Six Sigma Black Belt",
        "add_cert_button": "Add certification",
        "select_all_certs": "Select all certifications",
        "clear_certs": "Clear certifications",
        "certs_multiselect_label": "Certifications to include",
        "certs_multiselect_help": "Pick only the credentials you want to highlight.",
        "skills_expander_title": "Skills by Block",
        "skills_expander_caption": "Use the controls below to add, remove, or reorder skills for each text block.",
        "skill_add_prompt": "Add custom skill for {label}",
        "skill_add_placeholder": "e.g., Commercial Strategy",
        "skill_add_button": "Add skill",
        "skills_select_all": "Select all",
        "skills_clear": "Clear",
        "skills_multiselect_label": "Skills to include ({label})",
        "skills_multiselect_help": "Pick only the skills you want to mention in this block.",
        "contact_expander": "Contact Details",
        "contact_label_header": "Label",
        "contact_value_header": "Value",
        "add_contact_button": "Add contact",
        "reset_contacts_button": "Reset",
        "text_area_subheader": "Cover letter text",
        "text_area_label": "You can edit the text before downloading:",
        "preview_heading": "Preview (A4 layout)",
        "file_name_label": "File name (without extension)",
        "download_docx": "Download as Word (.docx)",
        "download_pdf": "Download as PDF (.pdf)",
        "footer_tip": (
            "Tip: Update 'profile.yaml' for reusable info like contact links, tweak 'degrees.yaml' / "
            "'certifications.yaml' for credentials, and edit the 'skills_*.yaml' files to set default skills per block."
        ),
    },
    "de": {
        "language_select_label": "Sprache",
        "candidate_expander": "Bewerberdetails",
        "full_name_label": "Vollständiger Name",
        "full_name_placeholder": "Vollständiger Name",
        "address_label": "Adresse",
        "address_placeholder": "Adresse",
        "city_label": "Stadt",
        "city_placeholder": "Stadt",
        "region_label": "Bundesland",
        "region_placeholder": "Bundesland",
        "country_label": "Land",
        "country_placeholder": "Land",
        "company_expander": "Unternehmensdetails",
        "company_name_label": "Unternehmensname",
        "company_name_placeholder": "Unternehmensname",
        "position_label": "Jobtitel / Position",
        "position_placeholder": "Jobtitel",
        "hiring_manager_label": "Ansprechperson",
        "hiring_manager_placeholder": "Ansprechperson",
        "job_city_label": "Stadt",
        "job_city_placeholder": "Stadt",
        "job_region_label": "Region",
        "job_region_placeholder": "Region/Bundesland",
        "job_country_label": "Land",
        "job_country_placeholder": "Land",
        "awareness_label": "Wie haben Sie von der Stelle erfahren?",
        "awareness_placeholder": "einen LinkedIn-Beitrag",
        "sections_expander": "Textbausteine",
        "section_candidate_block": "Bewerberblock",
        "section_company_block": "Unternehmensblock",
        "section_subject_block": "Betreffzeile",
        "section_greeting_block": "Begrüßung",
        "section_intro": "Intro & Fähigkeiten",
        "section_job_a": "Jobblock A",
        "section_job_b": "Jobblock B",
        "section_job_c": "Jobblock C",
        "section_job_d": "Jobblock D",
        "section_job_e": "Jobblock E",
        "section_degrees": "Abschlüsse & Zertifizierungen",
        "section_languages": "Sprachen & warum dieses Unternehmen",
        "section_goodbye": "Abschlussblock",
        "qualifications_expander": "Abschlüsse und Zertifizierungen",
        "qualifications_caption": "Wählen Sie, welche Abschlüsse und Zertifizierungen erscheinen oder fügen Sie neue hinzu.",
        "degrees_heading": "Abschlüsse",
        "degrees_input_label": "Eigenen Abschluss hinzufügen",
        "add_degree_placeholder": "z. B. MSc Big Data & Analytics",
        "add_degree_button": "Abschluss hinzufügen",
        "select_all_degrees": "Alle Abschlüsse wählen",
        "clear_degrees": "Abschlüsse leeren",
        "degrees_multiselect_label": "Anzuzeigende Abschlüsse",
        "degrees_multiselect_help": "Diese Angaben speisen den Satz „Ich verfüge über ...“.",
        "certifications_heading": "Zertifizierungen",
        "certs_input_label": "Eigene Zertifizierung hinzufügen",
        "add_cert_placeholder": "z. B. Six Sigma Black Belt",
        "add_cert_button": "Zertifizierung hinzufügen",
        "select_all_certs": "Alle Zertifizierungen wählen",
        "clear_certs": "Zertifizierungen leeren",
        "certs_multiselect_label": "Zertifizierungen auswählen",
        "certs_multiselect_help": "Markieren Sie nur die Zertifizierungen, die Sie hervorheben möchten.",
        "skills_expander_title": "Fähigkeiten je Abschnitt",
        "skills_expander_caption": "Fügen Sie pro Textblock Fähigkeiten hinzu oder entfernen Sie sie.",
        "skill_add_prompt": "Eigene Fähigkeit für {label} hinzufügen",
        "skill_add_placeholder": "z. B. Kommerzielle Strategie",
        "skill_add_button": "Fähigkeit hinzufügen",
        "skills_select_all": "Alle wählen",
        "skills_clear": "Leeren",
        "skills_multiselect_label": "Fähigkeiten für {label}",
        "skills_multiselect_help": "Wählen Sie die Fähigkeiten, die im Abschnitt erscheinen sollen.",
        "contact_expander": "Kontaktangaben",
        "contact_label_header": "Bezeichnung",
        "contact_value_header": "Wert",
        "add_contact_button": "Kontakt hinzufügen",
        "reset_contacts_button": "Zurücksetzen",
        "text_area_subheader": "Anschreiben-Text",
        "text_area_label": "Bearbeiten Sie den Text vor dem Download:",
        "preview_heading": "Vorschau (A4-Layout)",
        "file_name_label": "Dateiname (ohne Endung)",
        "download_docx": "Als Word (.docx) herunterladen",
        "download_pdf": "Als PDF (.pdf) herunterladen",
        "footer_tip": (
            "Tipp: Aktualisieren Sie 'profile.yaml' für wiederkehrende Infos, passen Sie 'degrees.yaml' / "
            "'certifications.yaml' für Abschlüsse/Zertifizierungen an und pflegen Sie die Dateien 'skills_*.yaml' pro Block."
        ),
    },
}
DEFAULT_LANGUAGE_STRINGS = LANGUAGE_STRINGS[DEFAULT_LANGUAGE]


# ----------------- Helpers ----------------- #

def _load_yaml(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        data = yaml.safe_load(f) or {}
    if not isinstance(data, dict):
        return {}
    return data


def _load_list_from_file(path: Path, key: str) -> list[str] | None:
    if not path.exists():
        return None
    data = _load_yaml(path)
    values = data.get(key, [])
    if isinstance(values, list):
        return values
    return None


def _load_skill_sources(language: str) -> dict[str, list[str]]:
    config = LANGUAGE_OPTIONS.get(language, LANGUAGE_OPTIONS[DEFAULT_LANGUAGE])
    suffix = config.get("skill_suffix", "")
    sources: dict[str, list[str]] = {}
    for block in SKILL_BLOCKS:
        path = Path(f"{block['file_stem']}{suffix}.yaml")
        sources[block["id"]] = _load_list_from_file(path, "skills") or []
    return sources


def load_profile():
    profile = _load_yaml(PROFILE_FILE)

    degrees = _load_list_from_file(DEGREES_FILE, "degrees")
    certifications = _load_list_from_file(CERTIFICATIONS_FILE, "certifications")

    profile["degrees"] = degrees if degrees is not None else profile.get("degrees", []) or []
    profile["certifications"] = (
        certifications if certifications is not None else profile.get("certifications", []) or []
    )
    return profile


def load_template(template_path: Path):
    with template_path.open("r", encoding="utf-8") as f:
        return Template(f.read(), trim_blocks=True, lstrip_blocks=True)


def build_candidate_lines(context: dict[str, str], sections_state: dict[str, bool]) -> list[str]:
    lines: list[str] = []
    if not sections_state.get("candidate_block"):
        return lines
    for key in ("name", "address_line1", "address_line2"):
        value = context.get(key)
        if value:
            lines.append(value.strip())
    home_location = context.get("home_location")
    if home_location:
        lines.append(home_location)
    return lines


def extract_candidate_lines_from_text(text: str, default_lines: list[str]) -> tuple[list[str], int]:
    if not default_lines:
        return [], 0
    lines = text.split("\n")
    collected: list[str] = []
    consumed = 0
    started = False
    for line in lines:
        if not started and not line.strip():
            consumed += 1
            continue
        if not started:
            started = True
        if line.startswith(" "):
            collected.append(line.strip())
            consumed += 1
            continue
        if not line.strip() and collected:
            consumed += 1
            continue
        break
    if collected:
        return collected, consumed
    return default_lines, len(default_lines)



def _split_blocks(text: str, skip_lines: int = 0) -> list[list[str]]:
    blocks: list[list[str]] = []
    lines = text.split("\n")
    remaining = lines[skip_lines:]
    while remaining and not remaining[0].strip():
        remaining.pop(0)
    current: list[str] = []
    for line in remaining:
        if not line.strip():
            if current:
                blocks.append(current)
                current = []
            continue
        current.append(line.rstrip())
    if current:
        blocks.append(current)
    return blocks


def assemble_text_value(
    candidate_lines: list[str], blocks: list[list[str]], include_company_block: bool
) -> str:
    """Rebuild editable text with consistent spacing."""

    text_blocks: list[str] = []
    block_index = 0
    if candidate_lines:
        if include_company_block and blocks:
            combined = candidate_lines + blocks[0]
            text_blocks.append("\n".join(combined))
            block_index = 1
        else:
            text_blocks.append("\n".join(candidate_lines))

    for block in blocks[block_index:]:
        block_text = "\n".join(block).strip("\n")
        if block_text:
            text_blocks.append(block_text)

    return "\n\n".join(text_blocks).strip()


def build_docx(
    text: str,
    sections_state: dict[str, bool],
    candidate_lines: list[str],
    candidate_line_consumed: int,
) -> BytesIO:
    """
    Create an in-memory .docx file from the cover letter text.
    """
    doc = Document()
    candidate_count = len(candidate_lines) if sections_state.get("candidate_block") else 0
    blocks = _split_blocks(text, skip_lines=candidate_line_consumed if candidate_count else 0)

    def add_block(lines: list[str], align_right: bool = False) -> None:
        for line in lines:
            paragraph = doc.add_paragraph(line.strip() if align_right else line)
            if align_right:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.line_spacing = 1

    if candidate_count:
        add_block(candidate_lines, align_right=True)

    for idx, block in enumerate(blocks):
        add_block(block)
        if idx < len(blocks) - 1:
            spacer = doc.add_paragraph("")
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.space_before = Pt(0)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_pdf(
    text: str,
    sections_state: dict[str, bool],
    candidate_lines: list[str],
    candidate_line_consumed: int,
) -> BytesIO:
    """Create an in-memory PDF using ReportLab platypus for better layout control."""

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=72,
        rightMargin=72,
        topMargin=72,
        bottomMargin=72,
    )

    candidate_count = len(candidate_lines) if sections_state.get("candidate_block") else 0
    blocks = _split_blocks(text, skip_lines=candidate_line_consumed if candidate_count else 0)

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=12,
        leading=14,
        alignment=TA_LEFT,
        spaceBefore=0,
        spaceAfter=8,
    )
    candidate_style = ParagraphStyle(
        "Candidate",
        parent=body_style,
        alignment=TA_RIGHT,
        spaceAfter=6,
    )

    story: list = []

    def para_text(lines: list[str]) -> str:
        return "<br/>".join(html.escape(line) if line.strip() else "&nbsp;" for line in lines)

    if candidate_count:
        story.append(Paragraph(para_text(candidate_lines), candidate_style))

    for idx, block in enumerate(blocks):
        story.append(Paragraph(para_text(block), body_style))

    if not story:
        story.append(Paragraph("", body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer


def _render_skill_selector(
    block_id: str,
    label: str,
    base_options: list[str],
    language_code: str,
    translate: Callable[..., str],
) -> list[str]:
    """Render a skills multiselect with add/remove helpers for a specific block."""
    st.markdown(f"**{label}**")
    key_prefix = f"{language_code}_{block_id}"
    custom_key = f"{key_prefix}_custom_skills"
    selection_key = f"{key_prefix}_skills_selection"
    input_key = f"{key_prefix}_skill_input"
    add_button_key = f"{key_prefix}_add_skill_button"
    select_all_key = f"{key_prefix}_select_all_skills"
    clear_key = f"{key_prefix}_clear_skills"

    st.session_state.setdefault(custom_key, [])

    new_skill = st.text_input(
        translate("skill_add_prompt", label=label),
        key=input_key,
        placeholder=translate("skill_add_placeholder"),
    )
    if st.button(f"➕ {translate('skill_add_button')}", key=add_button_key):
        value = new_skill.strip()
        if value:
            custom_entries = st.session_state[custom_key]
            if value not in base_options and value not in custom_entries:
                custom_entries.append(value)
                st.session_state[custom_key] = custom_entries
            current_selection = st.session_state.get(selection_key, [])
            if value not in current_selection:
                st.session_state[selection_key] = current_selection + [value]
        st.session_state[input_key] = ""
        st.experimental_rerun()

    options = base_options + [entry for entry in st.session_state[custom_key] if entry not in base_options]

    if selection_key not in st.session_state:
        st.session_state[selection_key] = list(options)
    else:
        st.session_state[selection_key] = [item for item in st.session_state[selection_key] if item in options]

    btn_cols = st.columns(2)
    if btn_cols[0].button(translate("skills_select_all"), key=select_all_key):
        st.session_state[selection_key] = list(options)
    if btn_cols[1].button(translate("skills_clear"), key=clear_key):
        st.session_state[selection_key] = []

    st.multiselect(
        translate("skills_multiselect_label", label=label),
        options=options,
        key=selection_key,
        help=translate("skills_multiselect_help"),
    )
    return st.session_state.get(selection_key, [])


# ----------------- Streamlit UI ----------------- #

st.set_page_config(page_title="Cover Letter Lego Builder", layout="wide")

st.title("Cover Letter Lego Builder")

profile = load_profile()

language_codes = list(LANGUAGE_OPTIONS.keys())
default_index = language_codes.index(DEFAULT_LANGUAGE) if DEFAULT_LANGUAGE in language_codes else 0
selected_language = st.sidebar.selectbox(
    LANGUAGE_STRINGS[DEFAULT_LANGUAGE]["language_select_label"],
    options=language_codes,
    index=default_index,
    format_func=lambda code: LANGUAGE_OPTIONS[code]["label"],
    key="language_select",
)
language_config = LANGUAGE_OPTIONS.get(selected_language, LANGUAGE_OPTIONS[DEFAULT_LANGUAGE])
lang_strings = LANGUAGE_STRINGS.get(selected_language, DEFAULT_LANGUAGE_STRINGS)

def t(key: str, **kwargs: str) -> str:
    value = lang_strings.get(key, DEFAULT_LANGUAGE_STRINGS.get(key, key))
    if kwargs:
        try:
            return value.format(**kwargs)
        except Exception:
            return value
    return value

template = load_template(language_config["template_path"])
skills_sources = _load_skill_sources(selected_language)

JOB_FIELDS = ("company", "position", "hiring_manager", "city", "region", "country")
for field in JOB_FIELDS:
    st.session_state.setdefault(f"{field}_input", profile.get(field, ""))

candidate_expander = st.sidebar.expander(t("candidate_expander"), expanded=False)
with candidate_expander:
    full_name = st.text_input(
        t("full_name_label"),
        value=profile.get("name", t("full_name_placeholder")),
        placeholder=t("full_name_placeholder"),
        key="candidate_full_name",
    )
    candidate_address = st.text_input(
        t("address_label"),
        value=profile.get("address_line1", t("address_placeholder")),
        placeholder=t("address_placeholder"),
        key="candidate_address",
    )
    candidate_city = st.text_input(
        t("city_label"),
        value=profile.get("city", t("city_placeholder")),
        placeholder=t("city_placeholder"),
        key="candidate_city",
    )
    candidate_region = st.text_input(
        t("region_label"),
        value=profile.get("region", t("region_placeholder")),
        placeholder=t("region_placeholder"),
        key="candidate_region",
    )
    candidate_country = st.text_input(
        t("country_label"),
        value=profile.get("country", t("country_placeholder")),
        placeholder=t("country_placeholder"),
        key="candidate_country",
    )

home_city = candidate_city.strip() or profile.get("city", "")
home_region = candidate_region.strip() or profile.get("region", "")
home_country = candidate_country.strip() or profile.get("country", "")
home_location = ", ".join([part for part in (home_city, home_region, home_country) if part])

company_expander = st.sidebar.expander(t("company_expander"), expanded=False)
with company_expander:
    company_value = st.text_input(
        t("company_name_label"),
        value=t("company_name_placeholder"),
        placeholder=t("company_name_placeholder"),
        key="company_name",
    )
    position_value = st.text_input(
        t("position_label"),
        value=t("position_placeholder"),
        placeholder=t("position_placeholder"),
        key="company_position",
    )
    hiring_manager_value = st.text_input(
        t("hiring_manager_label"),
        value=t("hiring_manager_placeholder"),
        placeholder=t("hiring_manager_placeholder"),
        key="company_hiring_manager",
    )
    city_value = st.text_input(
        t("job_city_label"),
        value=t("job_city_placeholder"),
        placeholder=t("job_city_placeholder"),
        key="company_city",
    )
    region_value = st.text_input(
        t("job_region_label"),
        value=t("job_region_placeholder"),
        placeholder=t("job_region_placeholder"),
        key="company_region",
    )
    country_value = st.text_input(
        t("job_country_label"),
        value=t("job_country_placeholder"),
        placeholder=t("job_country_placeholder"),
        key="company_country",
    )
    awareness = st.text_input(
        t("awareness_label"),
        value=t("awareness_placeholder"),
        placeholder=t("awareness_placeholder"),
        key="company_awareness",
    )

sections_expander = st.sidebar.expander(t("sections_expander"), expanded=False)
with sections_expander:
    sections = {
        "candidate_block": st.checkbox(t("section_candidate_block"), True),
        "company_block": st.checkbox(t("section_company_block"), True),
        "subject_block": st.checkbox(t("section_subject_block"), True),
        "greeting_block": st.checkbox(t("section_greeting_block"), True),
        "intro": st.checkbox(t("section_intro"), True),
        "market_dev": st.checkbox(t("section_job_a"), True),
        "senior_bi": st.checkbox(t("section_job_b"), True),
        "project_manager": st.checkbox(t("section_job_c"), True),
        "business_analytics": st.checkbox(t("section_job_d"), True),
        "fpna": st.checkbox(t("section_job_e"), True),
        "degrees_certs": st.checkbox(t("section_degrees"), True),
        "languages_company": st.checkbox(t("section_languages"), True),
        "goodbye_block": st.checkbox(t("section_goodbye"), True),
    }

degrees_source = profile.get("degrees", [])
certifications_source = profile.get("certifications", [])

qualifications_expander = st.sidebar.expander(t("qualifications_expander"), expanded=False)
with qualifications_expander:
    st.caption(t("qualifications_caption"))

    # Degrees management
    st.session_state.setdefault("custom_degree_entries", [])
    st.markdown(f"**{t('degrees_heading')}**")
    new_degree = st.text_input(
        t("degrees_input_label"),
        key="degree_add_input",
        placeholder=t("add_degree_placeholder"),
    )
    if st.button(f"➕ {t('add_degree_button')}", key="add_degree_button"):
        value = new_degree.strip()
        if value:
            if value not in degrees_source and value not in st.session_state["custom_degree_entries"]:
                st.session_state["custom_degree_entries"].append(value)
            current_selection = st.session_state.get("degrees_selection", [])
            if value not in current_selection:
                st.session_state["degrees_selection"] = current_selection + [value]
        st.session_state.degree_add_input = ""
        st.experimental_rerun()
    degree_options = degrees_source + st.session_state["custom_degree_entries"]
    if "degrees_selection" not in st.session_state:
        st.session_state["degrees_selection"] = list(degree_options)
    else:
        st.session_state["degrees_selection"] = [
            degree for degree in st.session_state["degrees_selection"] if degree in degree_options
        ]
    degree_button_cols = st.columns(2)
    if degree_button_cols[0].button(t("select_all_degrees"), key="select_all_degrees_button"):
        st.session_state["degrees_selection"] = list(degree_options)
    if degree_button_cols[1].button(t("clear_degrees"), key="clear_degrees_button"):
        st.session_state["degrees_selection"] = []
    st.multiselect(
        t("degrees_multiselect_label"),
        options=degree_options,
        key="degrees_selection",
        help=t("degrees_multiselect_help"),
    )

    st.markdown("---")

    # Certifications management
    st.session_state.setdefault("custom_cert_entries", [])
    st.markdown(f"**{t('certifications_heading')}**")
    new_cert = st.text_input(
        t("certs_input_label"),
        key="cert_add_input",
        placeholder=t("add_cert_placeholder"),
    )
    if st.button(f"➕ {t('add_cert_button')}", key="add_cert_button"):
        value = new_cert.strip()
        if value:
            if value not in certifications_source and value not in st.session_state["custom_cert_entries"]:
                st.session_state["custom_cert_entries"].append(value)
            current_certs = st.session_state.get("certifications_selection", [])
            if value not in current_certs:
                st.session_state["certifications_selection"] = current_certs + [value]
        st.session_state.cert_add_input = ""
        st.experimental_rerun()
    cert_options = certifications_source + st.session_state["custom_cert_entries"]
    if "certifications_selection" not in st.session_state:
        st.session_state["certifications_selection"] = list(cert_options)
    else:
        st.session_state["certifications_selection"] = [
            cert for cert in st.session_state["certifications_selection"] if cert in cert_options
        ]
    cert_button_cols = st.columns(2)
    if cert_button_cols[0].button(t("select_all_certs"), key="select_all_cert_button"):
        st.session_state["certifications_selection"] = list(cert_options)
    if cert_button_cols[1].button(t("clear_certs"), key="clear_cert_button"):
        st.session_state["certifications_selection"] = []
    st.multiselect(
        t("certs_multiselect_label"),
        options=cert_options,
        key="certifications_selection",
        help=t("certs_multiselect_help"),
    )

skills_expander = st.sidebar.expander(t("skills_expander_title"), expanded=False)
selected_skills_map: dict[str, list[str]] = {}
with skills_expander:
    st.caption(t("skills_expander_caption"))
    for idx, block in enumerate(SKILL_BLOCKS):
        block_id = block["id"]
        base_options = skills_sources.get(block_id, [])
        block_label = t(block.get("label_key", block_id))
        selected = _render_skill_selector(block_id, block_label, base_options, selected_language, t)
        selected_skills_map[block["context_key"]] = selected
        if idx < len(SKILL_BLOCKS) - 1:
            st.markdown("---")

selected_degrees = st.session_state.get("degrees_selection", degrees_source) or []
selected_certifications = st.session_state.get("certifications_selection", certifications_source) or []

contact_defaults = profile.get("contact_links", [])

def _sanitize_contact(entry: dict[str, str] | str, fallback_label: str) -> dict[str, str]:
    if isinstance(entry, dict):
        return {
            "label": entry.get("label") or fallback_label,
            "value": entry.get("value", ""),
        }
    if isinstance(entry, str):
        return {"label": fallback_label, "value": entry}
    return {"label": fallback_label, "value": ""}


def _build_default_contacts(raw_defaults: list) -> list[dict[str, str]]:
    fallback_labels = ["Email", "LinkedIn", "Xing", "Website"]
    defaults: list[dict[str, str]] = []
    if not raw_defaults:
        return defaults
    for idx, item in enumerate(raw_defaults):
        fallback = fallback_labels[idx] if idx < len(fallback_labels) else f"Contact {idx + 1}"
        defaults.append(_sanitize_contact(item, fallback))
    return defaults

if "contact_entries" not in st.session_state:
    defaults = _build_default_contacts(contact_defaults)
    if defaults:
        st.session_state.contact_entries = defaults
    else:
        st.session_state.contact_entries = [
            {"label": "Email", "value": profile.get("applicant", {}).get("email", "")}
        ]

custom_expander = st.sidebar.expander(t("contact_expander"), expanded=False)
with custom_expander:
    remove_indices: list[int] = []
    header_cols = st.columns([1.1, 2.2, 0.4])
    header_cols[0].markdown(f"**{t('contact_label_header')}**")
    header_cols[1].markdown(f"**{t('contact_value_header')}**")
    header_cols[2].markdown(" ")

    for idx, entry in enumerate(st.session_state.contact_entries):
        cols = st.columns([1.1, 2.2, 0.4])
        label_key = f"contact_label_{idx}"
        value_key = f"contact_value_{idx}"
        label_value = cols[0].text_input(
            "Label",
            entry["label"],
            key=label_key,
            label_visibility="collapsed",
        )
        value_value = cols[1].text_input(
            "Value",
            entry["value"],
            key=value_key,
            label_visibility="collapsed",
        )
        st.session_state.contact_entries[idx] = {
            "label": label_value.strip() or "Label",
            "value": value_value.strip(),
        }
        if cols[2].button("🗑️", key=f"contact_remove_{idx}") and len(st.session_state.contact_entries) > 1:
            remove_indices.append(idx)

    if remove_indices:
        for idx in sorted(remove_indices, reverse=True):
            st.session_state.contact_entries.pop(idx)

    add_col1, add_col2 = st.columns([1, 1])
    if add_col1.button(f"➕ {t('add_contact_button')}"):
        st.session_state.contact_entries.append({"label": f"Contact {len(st.session_state.contact_entries) + 1}", "value": ""})
    if add_col2.button(f"↻ {t('reset_contacts_button')}"):
        defaults = _build_default_contacts(contact_defaults)
        st.session_state.contact_entries = defaults or [
            {"label": "Email", "value": profile.get("applicant", {}).get("email", "")}
        ]

contact_links = [
    entry
    for entry in st.session_state.contact_entries
    if entry.get("value")
]

# Use sidebar inputs for job information
job_info = {
    "company": company_value.strip(),
    "position": position_value.strip(),
    "hiring_manager": hiring_manager_value.strip(),
    "city": city_value.strip(),
    "region": region_value.strip(),
    "country": country_value.strip(),
}

profile_for_context = {**profile}
profile_for_context.update(
    {
        "name": full_name.strip() or profile.get("name", ""),
        "address_line1": candidate_address.strip() or profile.get("address_line1", ""),
        "city": home_city,
        "region": home_region,
        "country": home_country,
    }
)
profile_for_context["degrees"] = selected_degrees
profile_for_context["certifications"] = selected_certifications
for block in SKILL_BLOCKS:
    ctx_key = block["context_key"]
    profile_for_context[ctx_key] = selected_skills_map.get(ctx_key, skills_sources.get(block["id"], []))

# Merge everything into a single context for the template
context = {
    **profile_for_context,
    **job_info,
    "hiring_manager": job_info.get("hiring_manager") or None,
    "awareness_source": awareness,
    "sections": sections,
    "home_city": home_city,
    "home_region": home_region,
    "home_country": home_country,
    "home_location": home_location,
    "contact_links": contact_links or profile.get("contact_links", []),
}

candidate_lines_plain = build_candidate_lines(context, sections)

# Render the cover letter
rendered_text = template.render(**context)
if sections.get("candidate_block") and sections.get("company_block"):
    rendered_text = rendered_text.replace("\n\n", "\n", 1)

initial_candidate_lines, initial_consumed = extract_candidate_lines_from_text(
    rendered_text, candidate_lines_plain
)
initial_blocks = _split_blocks(
    rendered_text,
    skip_lines=initial_consumed if initial_candidate_lines else 0,
)

text_area_value = assemble_text_value(
    initial_candidate_lines,
    initial_blocks,
    bool(sections.get("company_block")),
)
if not text_area_value:
    text_area_value = rendered_text

st.subheader(t("text_area_subheader"))
edited_text = st.text_area(
    t("text_area_label"),
    value=text_area_value,
    height=450,
)

preview_candidate_lines, preview_consumed = extract_candidate_lines_from_text(
    edited_text, candidate_lines_plain
)
preview_blocks = _split_blocks(
    edited_text,
    skip_lines=preview_consumed if preview_candidate_lines else 0,
)

preview_parts = []
if preview_candidate_lines:
    preview_parts.append(
        "<div style='text-align:right;'>" + "<br>".join(html.escape(line) for line in preview_candidate_lines) + "</div>"
    )
for block in preview_blocks:
    preview_parts.append(
        "<div style='text-align:left;'>"
        + "<br>".join(html.escape(line) for line in block)
        + "</div>"
    )
preview_html = (
    "<div style='font-family: "
    "Courier New, monospace; white-space: pre-wrap; border:1px solid #ddd; padding:1rem; background:#f8f9fb;'>"
    + "<div style='display:flex; flex-direction:column; gap:1rem;'>"
    + "".join(preview_parts)
    + "</div></div>"
)

st.markdown(f"**{t('preview_heading')}**")
st.markdown(preview_html, unsafe_allow_html=True)

default_export_name = " - ".join(
    filter(
        None,
        [
            company_value.strip(),
            position_value.strip(),
            context.get("today"),
        ],
    )
) or "cover_letter"

state = st.session_state
if "export_filename_state" not in state:
    state.export_filename_state = default_export_name
    state.export_name_custom = False

if not state.get("export_name_custom", False):
    state.export_filename_state = default_export_name

export_name_input = st.text_input(
    t("file_name_label"),
    value=state.export_filename_state,
    key="export_filename_input",
)
raw_export_name = export_name_input.strip()
if not raw_export_name:
    export_name = default_export_name
    state.export_filename_state = default_export_name
    state.export_name_custom = False
else:
    export_name = raw_export_name
    state.export_filename_state = export_name
    state.export_name_custom = export_name != default_export_name

candidate_lines_for_export, consumed_candidate_lines = extract_candidate_lines_from_text(
    edited_text, candidate_lines_plain
)

col1, col2 = st.columns(2)

with col1:
    docx_file = build_docx(
        edited_text,
        sections,
        candidate_lines_for_export,
        consumed_candidate_lines,
    )
    st.download_button(
        t("download_docx"),
        data=docx_file,
        file_name=f"{export_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

with col2:
    pdf_file = build_pdf(
        edited_text,
        sections,
        candidate_lines_for_export,
        consumed_candidate_lines,
    )
    st.download_button(
        t("download_pdf"),
        data=pdf_file,
        file_name=f"{export_name}.pdf",
        mime="application/pdf",
    )

st.caption(t("footer_tip"))
